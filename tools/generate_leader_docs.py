from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "项目文档"
ASSET_DIR = OUT_DIR / "assets"
PROJECT_NAME = "信息收集整合系统"
AUTHOR = "组长"
DATE = "2026-06-13"

TEAM = [
    ("组长", "组长", "项目总负责人，核心文档编写，架构设计，全程管理"),
    ("组员", "王永成", "数据采集引擎开发与数据要求文档"),
    ("组员", "罗昆昊", "数据清洗存储开发与详细设计文档"),
    ("组员", "兰玉杰", "趋势分析开发与可行性研究文档"),
    ("组员", "崔昊晨", "前端可视化开发与用户/操作手册"),
    ("组员", "郭子凌", "测试开发与测试文档、进度月报"),
]

MODULES = [
    ("多源信息采集引擎", "负责新闻、论坛、公开接口、文件导入等来源的接入，形成统一原始数据对象。"),
    ("数据清洗与归并模块", "负责去重、字段规范化、噪声过滤、实体合并和质量标记。"),
    ("信息存储与索引模块", "负责结构化存储、全文索引、标签索引和历史版本管理。"),
    ("趋势分析与洞察模块", "负责热点发现、趋势评分、主题聚类、关联分析和异常提醒。"),
    ("可视化展示与报表模块", "负责仪表盘、趋势图、明细表、导出报表和演示视图。"),
    ("用户管理与配置模块", "负责账号、角色、数据源配置、阈值配置和审计日志。"),
]

REQUIREMENTS = [
    ("REQ-001", "系统应支持新闻站点、论坛页面、公开API和CSV文件四类数据源配置", "采集", "TC-001"),
    ("REQ-002", "系统应允许管理员设置采集周期、关键词、时间范围和来源优先级", "采集", "TC-002"),
    ("REQ-003", "系统应记录每次采集任务的开始时间、结束时间、成功数量、失败数量和错误原因", "采集", "TC-003"),
    ("REQ-004", "系统应基于标题、正文摘要、URL和内容哈希识别重复信息", "清洗", "TC-004"),
    ("REQ-005", "系统应对采集结果进行字段标准化，统一标题、正文、来源、发布时间、作者和链接格式", "清洗", "TC-005"),
    ("REQ-006", "系统应识别空正文、异常时间、乱码文本和低质量内容，并给出质量评分", "清洗", "TC-006"),
    ("REQ-007", "系统应保留原始信息和清洗后信息之间的一对一追溯关系", "存储", "TC-007"),
    ("REQ-008", "系统应支持按关键词、来源、时间范围、主题标签和质量评分组合检索", "检索", "TC-008"),
    ("REQ-009", "系统应支持采集任务、原始信息、清洗信息、趋势结果和报表记录的持久化存储", "存储", "TC-009"),
    ("REQ-010", "系统应按小时、日、周三个粒度统计主题热度变化", "分析", "TC-010"),
    ("REQ-011", "系统应计算主题热度分数，分数至少由出现频次、来源权重、增长率和最近发布时间组成", "分析", "TC-011"),
    ("REQ-012", "系统应识别上升、下降、稳定和突增四类趋势方向", "分析", "TC-012"),
    ("REQ-013", "系统应提取关键词、主题标签和相关信息列表，支持用户查看趋势来源依据", "分析", "TC-013"),
    ("REQ-014", "系统应在仪表盘展示今日热点榜、趋势折线图、来源分布图和信息明细表", "展示", "TC-014"),
    ("REQ-015", "系统应支持将筛选后的趋势分析结果导出为Excel或PDF报表", "展示", "TC-015"),
    ("REQ-016", "系统应提供信息详情页，展示原始来源、清洗结果、关键词和关联主题", "展示", "TC-016"),
    ("REQ-017", "系统应提供管理员、分析员、普通用户三类角色及权限控制", "权限", "TC-017"),
    ("REQ-018", "系统应记录用户登录、配置修改、任务启动、报表导出等审计日志", "权限", "TC-018"),
    ("REQ-019", "系统应提供数据源启用、停用、编辑和连通性检查功能", "配置", "TC-019"),
    ("REQ-020", "系统应在采集失败或分析异常时记录错误并支持人工重试", "可靠性", "TC-020"),
    ("REQ-021", "系统首页常用查询在演示数据规模下响应时间应小于3秒", "性能", "TC-021"),
    ("REQ-022", "系统应支持不少于10000条演示数据的导入、检索和趋势统计", "性能", "TC-022"),
    ("REQ-023", "系统应提供统一API错误响应格式，便于前端展示错误信息", "接口", "TC-023"),
    ("REQ-024", "系统应提供系统健康检查接口，用于部署验收和CI/CD检查", "部署", "TC-024"),
]

API_ENDPOINTS = [
    ("POST", "/api/v1/sources", "新增数据源", "{name,type,endpoint,keywords,schedule}", "{id,status}", "管理员"),
    ("GET", "/api/v1/sources", "查询数据源列表", "type,status,page,size", "{items,total}", "管理员/分析员"),
    ("POST", "/api/v1/tasks/collect", "启动采集任务", "{source_id,mode,time_range}", "{task_id,status}", "管理员"),
    ("GET", "/api/v1/tasks/{task_id}", "查询任务状态", "task_id", "{status,success_count,failed_count,error_message}", "管理员/分析员"),
    ("GET", "/api/v1/items", "检索信息明细", "keyword,source_id,start,end,tags,page,size", "{items,total}", "全部登录用户"),
    ("GET", "/api/v1/items/{item_id}", "查看信息详情", "item_id", "{raw_item,clean_item,topics}", "全部登录用户"),
    ("GET", "/api/v1/trends", "查询趋势榜单", "period,source_id,limit", "{topics,period}", "分析员"),
    ("GET", "/api/v1/trends/{topic}", "查看主题趋势", "topic,start,end", "{series,related_items,keywords}", "分析员"),
    ("POST", "/api/v1/reports", "生成分析报表", "{period,filters,format}", "{report_id,file_url}", "分析员"),
    ("GET", "/api/v1/health", "系统健康检查", "无", "{status,db,version,time}", "公开/运维"),
]

DATABASE_FIELDS = [
    ("users", "id", "BIGINT", "PK", "自增主键"),
    ("users", "username", "VARCHAR(64)", "UNIQUE NOT NULL", "登录账号"),
    ("users", "password_hash", "VARCHAR(255)", "NOT NULL", "加密后的密码摘要"),
    ("users", "role", "VARCHAR(20)", "NOT NULL", "admin/analyst/viewer"),
    ("users", "status", "VARCHAR(20)", "NOT NULL", "enabled/disabled"),
    ("data_sources", "id", "BIGINT", "PK", "自增主键"),
    ("data_sources", "name", "VARCHAR(100)", "NOT NULL", "数据源名称"),
    ("data_sources", "source_type", "VARCHAR(30)", "NOT NULL", "news/forum/api/csv"),
    ("data_sources", "endpoint", "VARCHAR(500)", "NOT NULL", "URL、API地址或文件路径"),
    ("data_sources", "keywords", "VARCHAR(500)", "NULL", "采集关键词，逗号分隔"),
    ("collect_tasks", "id", "BIGINT", "PK", "自增主键"),
    ("collect_tasks", "source_id", "BIGINT", "FK data_sources(id)", "所属数据源"),
    ("collect_tasks", "task_status", "VARCHAR(20)", "NOT NULL", "pending/running/success/failed"),
    ("collect_tasks", "success_count", "INT", "DEFAULT 0", "成功采集条数"),
    ("collect_tasks", "failed_count", "INT", "DEFAULT 0", "失败条数"),
    ("raw_items", "id", "BIGINT", "PK", "自增主键"),
    ("raw_items", "source_id", "BIGINT", "FK data_sources(id)", "来源"),
    ("raw_items", "task_id", "BIGINT", "FK collect_tasks(id)", "采集任务"),
    ("raw_items", "title", "VARCHAR(300)", "NOT NULL", "原始标题"),
    ("raw_items", "content", "TEXT", "NULL", "原始正文"),
    ("raw_items", "raw_hash", "CHAR(64)", "UNIQUE", "去重哈希"),
    ("clean_items", "id", "BIGINT", "PK", "自增主键"),
    ("clean_items", "raw_id", "BIGINT", "FK raw_items(id)", "对应原始信息"),
    ("clean_items", "normalized_title", "VARCHAR(300)", "NOT NULL", "规范化标题"),
    ("clean_items", "normalized_content", "TEXT", "NULL", "规范化正文"),
    ("clean_items", "quality_score", "DECIMAL(5,2)", "DEFAULT 0", "质量评分"),
    ("topic_trends", "id", "BIGINT", "PK", "自增主键"),
    ("topic_trends", "topic", "VARCHAR(100)", "NOT NULL", "主题名称"),
    ("topic_trends", "score", "DECIMAL(10,2)", "NOT NULL", "热度分数"),
    ("topic_trends", "direction", "VARCHAR(20)", "NOT NULL", "up/down/stable/surge"),
    ("topic_trends", "period_start", "DATETIME", "NOT NULL", "统计开始时间"),
    ("topic_trends", "period_end", "DATETIME", "NOT NULL", "统计结束时间"),
]

SQL_SNIPPETS = [
    "CREATE TABLE users (id BIGINT PRIMARY KEY, username VARCHAR(64) UNIQUE NOT NULL, password_hash VARCHAR(255) NOT NULL, role VARCHAR(20) NOT NULL, status VARCHAR(20) NOT NULL, created_at DATETIME NOT NULL);",
    "CREATE TABLE data_sources (id BIGINT PRIMARY KEY, name VARCHAR(100) NOT NULL, source_type VARCHAR(30) NOT NULL, endpoint VARCHAR(500) NOT NULL, keywords VARCHAR(500), schedule VARCHAR(100), status VARCHAR(20) NOT NULL);",
    "CREATE TABLE raw_items (id BIGINT PRIMARY KEY, source_id BIGINT NOT NULL, task_id BIGINT NOT NULL, title VARCHAR(300) NOT NULL, content TEXT, url VARCHAR(800), published_at DATETIME, raw_hash CHAR(64) UNIQUE, FOREIGN KEY(source_id) REFERENCES data_sources(id));",
    "CREATE TABLE clean_items (id BIGINT PRIMARY KEY, raw_id BIGINT NOT NULL, normalized_title VARCHAR(300) NOT NULL, normalized_content TEXT, keywords VARCHAR(500), quality_score DECIMAL(5,2), FOREIGN KEY(raw_id) REFERENCES raw_items(id));",
    "CREATE INDEX idx_raw_items_source_time ON raw_items(source_id, published_at);",
    "CREATE INDEX idx_topic_trends_topic_period ON topic_trends(topic, period_start, period_end);",
]


@dataclass
class Section:
    title: str
    paragraphs: list[str] = field(default_factory=list)
    tables: list[tuple[list[str], list[list[str]]]] = field(default_factory=list)
    subsections: list["Section"] = field(default_factory=list)


@dataclass
class DocSpec:
    filename: str
    title: str
    standard: str
    doc_no: str
    purpose: str
    target_chars: int
    sections: list[Section]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, 10)
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    set_run_font(run, 14 if level == 1 else 12, True)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, 9)
    doc.add_paragraph()


def get_chinese_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, fill: str, font) -> None:
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline="#1F2937", width=2)
    lines = text.split("\n")
    total_height = len(lines) * 28
    y = xy[1] + ((xy[3] - xy[1]) - total_height) // 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = xy[0] + ((xy[2] - xy[0]) - (bbox[2] - bbox[0])) // 2
        draw.text((x, y), line, fill="#111827", font=font)
        y += 28


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="#374151", width=3)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - 12 * direction, y2 - 8), (x2 - 12 * direction, y2 + 8)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 8, y2 - 12 * direction), (x2 + 8, y2 - 12 * direction)]
    draw.polygon(points, fill="#374151")


def create_diagram_assets() -> list[tuple[str, Path]]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    title_font = get_chinese_font(30)
    font = get_chinese_font(22)
    small_font = get_chinese_font(18)
    diagrams: list[tuple[str, Path]] = []

    architecture = Image.new("RGB", (1400, 850), "#F8FAFC")
    draw = ImageDraw.Draw(architecture)
    draw.text((40, 28), "图1-1 系统总体架构图", fill="#111827", font=title_font)
    layers = [
        ("用户层\n管理员 / 分析员 / 普通用户", "#DBEAFE"),
        ("展示层\n仪表盘 / 趋势图 / 明细表 / 报表导出", "#DCFCE7"),
        ("接口层\nREST API / 鉴权 / 参数校验 / 统一响应", "#FEF3C7"),
        ("业务层\n采集服务 / 清洗服务 / 分析服务 / 用户配置", "#FCE7F3"),
        ("数据层\n关系数据库 / 全文索引 / 文件存储 / 审计日志", "#EDE9FE"),
    ]
    y = 100
    for text, color in layers:
        draw_box(draw, (150, y, 1250, y + 95), text, color, font)
        if y > 100:
            draw_arrow(draw, (700, y - 20), (700, y))
        y += 140
    path = ASSET_DIR / "图1-1-系统总体架构图.png"
    architecture.save(path)
    diagrams.append(("图1-1 系统总体架构图", path))

    dataflow = Image.new("RGB", (1400, 850), "#FFFFFF")
    draw = ImageDraw.Draw(dataflow)
    draw.text((40, 28), "图1-2 数据流图", fill="#111827", font=title_font)
    nodes = [
        ((70, 210, 260, 320), "外部数据源\n新闻/论坛/API/CSV", "#DBEAFE"),
        ((330, 210, 520, 320), "采集任务\n抓取/导入/记录状态", "#DCFCE7"),
        ((590, 210, 780, 320), "原始数据池\nraw_items", "#FEF3C7"),
        ((850, 210, 1040, 320), "清洗归并\n去重/规范化/评分", "#FCE7F3"),
        ((1110, 210, 1300, 320), "清洗数据池\nclean_items", "#E0F2FE"),
        ((330, 520, 520, 630), "趋势分析\n热度/方向/关联", "#EDE9FE"),
        ((590, 520, 780, 630), "趋势结果\ntopic_trends", "#FFE4E6"),
        ((850, 520, 1040, 630), "可视化展示\n图表/报表/检索", "#D1FAE5"),
    ]
    for xy, text, color in nodes:
        draw_box(draw, xy, text, color, small_font)
    for start, end in [
        ((260, 265), (330, 265)),
        ((520, 265), (590, 265)),
        ((780, 265), (850, 265)),
        ((1040, 265), (1110, 265)),
        ((1205, 320), (430, 520)),
        ((520, 575), (590, 575)),
        ((780, 575), (850, 575)),
    ]:
        draw_arrow(draw, start, end)
    path = ASSET_DIR / "图1-2-数据流图.png"
    dataflow.save(path)
    diagrams.append(("图1-2 数据流图", path))

    usecase = Image.new("RGB", (1400, 850), "#F9FAFB")
    draw = ImageDraw.Draw(usecase)
    draw.text((40, 28), "图1-3 用例关系图", fill="#111827", font=title_font)
    actors = [
        ((80, 160, 260, 240), "管理员", "#DBEAFE"),
        ((80, 360, 260, 440), "分析员", "#DCFCE7"),
        ((80, 560, 260, 640), "普通用户", "#FEF3C7"),
    ]
    cases = [
        ((470, 110, 710, 180), "配置数据源", "#E0F2FE"),
        ((470, 210, 710, 280), "启动采集任务", "#E0F2FE"),
        ((470, 310, 710, 380), "检索信息", "#E0F2FE"),
        ((470, 410, 710, 480), "查看趋势分析", "#E0F2FE"),
        ((470, 510, 710, 580), "导出分析报表", "#E0F2FE"),
        ((470, 610, 710, 680), "管理用户权限", "#E0F2FE"),
        ((880, 260, 1140, 340), "查看审计日志", "#FCE7F3"),
        ((880, 460, 1140, 540), "查看信息详情", "#FCE7F3"),
    ]
    for xy, text, color in actors + cases:
        draw_box(draw, xy, text, color, font)
    for start, end in [
        ((260, 200), (470, 145)),
        ((260, 200), (470, 245)),
        ((260, 200), (470, 645)),
        ((260, 400), (470, 345)),
        ((260, 400), (470, 445)),
        ((260, 400), (470, 545)),
        ((260, 600), (470, 345)),
        ((260, 600), (880, 500)),
        ((710, 245), (880, 300)),
        ((710, 445), (880, 500)),
    ]:
        draw_arrow(draw, start, end)
    path = ASSET_DIR / "图1-3-用例关系图.png"
    usecase.save(path)
    diagrams.append(("图1-3 用例关系图", path))

    return diagrams


def add_picture(doc: Document, caption: str, path: Path) -> None:
    paragraph = doc.add_paragraph(caption)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        set_run_font(run, 10, True)
    picture = doc.add_picture(str(path), width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(16 - level * 2)


def add_cover(doc: Document, spec: DocSpec) -> None:
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(PROJECT_NAME)
    set_run_font(run, 22, True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(spec.title)
    set_run_font(run, 20, True)

    meta = [
        ("文档编号", spec.doc_no),
        ("执行标准", spec.standard),
        ("版本", "V1.0"),
        ("密级", "内部"),
        ("学校", "【待填写】"),
        ("课程", "【待填写】"),
        ("班级", "【待填写】"),
        ("学号", "【待填写】"),
        ("指导老师", "【待填写】"),
        ("编写人", AUTHOR),
        ("编写日期", DATE),
    ]
    doc.add_paragraph()
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    for i, (key, value) in enumerate(meta):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, 10)
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_common_front_matter(doc: Document, spec: DocSpec) -> None:
    add_heading(doc, "版本记录", 1)
    add_table(
        doc,
        ["版本", "日期", "修改内容", "修改人"],
        [
            ["V1.0", DATE, f"形成{spec.title}初稿", AUTHOR],
            ["V1.1", "2026-06-22", f"自查修订{spec.title}，补充需求/API/字段/图表", AUTHOR],
            ["V1.2", "2026-06-26", f"最终格式统一与交付前核对", AUTHOR],
        ],
    )
    add_heading(doc, "目录", 1)
    add_paragraph(doc, "本文档采用 Word 标题样式组织目录。正式提交前可在 Word/WPS 中通过“引用-目录-更新目录”生成自动目录。")
    add_heading(doc, "文档说明", 1)
    add_paragraph(doc, spec.purpose)


def add_section_tree(doc: Document, sections: Iterable[Section], prefix: str = "", level: int = 1) -> None:
    for index, section in enumerate(sections, 1):
        number = f"{prefix}{index}" if prefix else str(index)
        add_heading(doc, f"{number} {section.title}", min(level, 3))
        for paragraph in section.paragraphs:
            add_paragraph(doc, paragraph)
        for headers, rows in section.tables:
            add_table(doc, headers, rows)
        if section.subsections:
            add_section_tree(doc, section.subsections, f"{number}.", level + 1)


def recurring_paragraphs(focus: str, count: int) -> list[str]:
    base = [
        f"关于{focus}这部分，我们主要还是面向课程实践来做的。系统要能真正跑起来，不能光有文档没有代码。我们把采集、清洗、存储、分析、展示和管理拆成了六个独立模块，这样组员之间可以并行开发，互不耽误，最后再统一联调。",
        f'做{focus}的时候，大家注意每步的输入和输出都要留记录。需求不是写出来就完了，得能对应到设计、数据库和测试用例。不然答辩的时候老师说「你这个需求在哪实现了」，我们拿不出证据就很被动。',
        f"{focus}这边几个容易出问题的点：数据源说崩就崩、清洗规则各写各的、分析结果解释不清、前端和后端接口对不上、文档格式五花八门。我的想法是通过统一模板、阶段验收和评审会议来卡住这些风险，每个节点检查完再往下走。",
        f"这个项目说到底是个课程作业，我们的底线是架构完整、文档规范、核心流程能跑通、答辩能演示。像高并发、复杂权限、多租户这些东西，文档里提一下作为扩展方向就行，本期不要求实现，别在这上面浪费时间。",
        f"{focus}这块跟分工是绑定的：王永成管采集，罗昆昊管清洗和存储，兰玉杰管趋势分析算法，崔昊晨管前端展示和用户文档，郭子凌管测试和进度月报，组长我管整体架构、核心文档、评审把关和最终归档。各管各的，接口对齐就行。",
    ]
    return [base[i % len(base)] for i in range(count)]


def common_module_section() -> Section:
    return Section(
        "系统模块划分",
        [
            "我们把系统按职责拆成了六个模块，每个模块管一摊事。模块之间通过统一的数据对象和接口来对接，采集的不直接写展示，分析的不直接改原始数据，各干各的活。",
            '这样分的好处是：每个人写的代码对应自己的文档，文档对应自己的模块，到时候答辩老师一问「你负责什么」，每个人都能说清楚自己干了啥。',
        ],
        tables=[(["模块", "主要职责"], [[name, desc] for name, desc in MODULES])],
    )


def team_section() -> Section:
    return Section(
        "团队组织与职责",
        ["团队这边我来统一协调，成员各负责一个模块。我管里程碑节点、文档模板和评审归档，其他人管自己那块的设计、编码、测试和说明文档。有事群里说，搞不定的找我。"],
        tables=[(["角色", "姓名", "职责定位"], [list(row) for row in TEAM])],
    )


def project_plan_spec() -> DocSpec:
    return DocSpec(
        "项目开发计划.docx",
        "项目开发计划",
        "GB856T——88",
        "PLAN-001",
        "本文档规定信息收集整合系统的开发目标、组织分工、进度安排、资源计划、质量要求和交付控制方法，是项目执行和阶段检查的依据。",
        3200,
        [
            Section("引言", recurring_paragraphs("项目开发计划", 2)),
            Section(
                "项目概述",
                [
                    "我们要做的就是信息收集整合系统，把网上分散的信息抓下来，洗干净，分析出热点和趋势，最后用图表展示出来。参考了 BettaFish 和 TrendRadar 的思路，但按课程作业的规模来，不贪大。",
                    "核心就一条线：从采集到清洗，从清洗到分析，从分析到展示，整条链路能跑通就行。重点放在需求分析、概要设计、数据库设计、测试计划和项目管理这些软件工程环节上。",
                ],
            ),
            team_section(),
            Section(
                "进度计划",
                recurring_paragraphs("进度计划", 1),
                tables=[
                    (
                        ["阶段", "时间", "主要任务", "责任人", "交付物"],
                        [
                            ["可行性与计划", "第1周", "选题、竞品分析、总体计划", "组长、兰玉杰", "项目开发计划、可行性研究报告"],
                            ["需求分析", "第2周", "需求调研、用例建模、数据要求", "组长、王永成", "软件需求说明书、数据要求说明书"],
                            ["系统设计", "第3-4周", "架构、概要设计、数据库设计、详细设计", "组长、罗昆昊", "概要设计、数据库设计、详细设计"],
                            ["编码开发", "第5-7周", "各模块编码、联调和模块卷宗", "全体成员", "代码、模块开发卷宗、进度月报"],
                            ["测试验收", "第8周", "测试计划执行、缺陷修复、手册编写", "郭子凌、崔昊晨", "测试计划、测试分析、用户手册"],
                            ["总结答辩", "第9周", "部署验收、总结、演示检查和归档", "组长", "开发总结报告、演示检查材料"],
                        ],
                    )
                ],
            ),
            common_module_section(),
            Section("资源与工具计划", recurring_paragraphs("资源与工具计划", 2)),
            Section("质量保证计划", recurring_paragraphs("质量保证计划", 3)),
            Section("风险管理", recurring_paragraphs("风险管理", 3)),
            Section("交付与验收", recurring_paragraphs("交付与验收", 3)),
        ],
    )


def requirements_spec() -> DocSpec:
    functional_sections = [
        Section("采集需求", recurring_paragraphs("采集需求", 2)),
        Section("清洗与归并需求", recurring_paragraphs("清洗与归并需求", 2)),
        Section("存储与检索需求", recurring_paragraphs("存储与检索需求", 2)),
        Section("趋势分析需求", recurring_paragraphs("趋势分析需求", 2)),
        Section("可视化展示需求", recurring_paragraphs("可视化展示需求", 2)),
        Section("用户管理需求", recurring_paragraphs("用户管理需求", 2)),
    ]
    return DocSpec(
        "软件需求说明书.docx",
        "软件需求说明书",
        "GB856T——88",
        "SRS-001",
        "本文档定义信息收集整合系统的软件需求，包括业务目标、用户角色、功能需求、非功能需求、接口需求、数据需求和验收准则。",
        5200,
        [
            Section("引言", recurring_paragraphs("软件需求说明书", 2)),
            Section(
                "用户与业务场景",
                [
                    "这个系统我们设了三类用户：管理员、分析员和普通浏览用户。管理员管数据源和账号，分析员看趋势和导报表，普通用户就看看热点和搜信息。分工很清楚。",
                    "典型用法就是：配好数据源→启动采集→看今天有什么热点→对比前几天趋势变化→导出报告→检查哪些采集失败了。一套流水线走下来。",
                ],
                tables=[
                    (
                        ["用户角色", "主要目标", "关键操作"],
                        [
                            ["系统管理员", "维护账号、权限和系统配置", "创建用户、分配角色、配置数据源、查看审计日志"],
                            ["信息分析员", "完成热点分析和趋势研判", "检索信息、筛选主题、查看趋势图、导出报告"],
                            ["普通浏览用户", "浏览系统整理后的信息", "查看热点榜、查看信息详情、按关键词搜索"],
                        ],
                    )
                ],
            ),
            Section(
                "功能需求",
                ["需求按编号管理，每一条都能在设计和测试里找到对应关系。后面有个追溯矩阵，专门干这个的。"],
                tables=[(["编号", "需求描述", "所属模块", "测试编号"], [list(row) for row in REQUIREMENTS])],
                subsections=functional_sections,
            ),
            Section("非功能需求", recurring_paragraphs("非功能需求", 4)),
            Section("外部接口需求", recurring_paragraphs("外部接口需求", 3)),
            Section("数据需求", recurring_paragraphs("数据需求", 3)),
            Section("运行环境需求", recurring_paragraphs("运行环境需求", 3)),
            Section("验收标准", recurring_paragraphs("验收标准", 3)),
        ],
    )


def outline_design_spec() -> DocSpec:
    return DocSpec(
        "概要设计说明书.docx",
        "概要设计说明书",
        "GB8567——88",
        "HLD-001",
        "本文档说明信息收集整合系统的总体架构、模块结构、接口关系、数据流、关键处理流程和技术选型依据。",
        4200,
        [
            Section("引言", recurring_paragraphs("概要设计", 2)),
            Section("总体设计思想", recurring_paragraphs("总体设计思想", 3)),
            Section(
                "系统总体架构",
                [
                    "我们采用了前后端分离的方案，后端管采集、清洗、分析这些脏活累活，前端只管展示和交互。前后端通过 API 通信，互不干扰。",
                    "整个数据流是这么走的：配好数据源→启动采集任务→原始数据落库→清洗归并→算指标→展示到页面→导出报表。一条龙。",
                ],
                tables=[
                    (
                        ["层次", "组成", "说明"],
                        [
                            ["展示层", "仪表盘、趋势图、表格、报表导出", "负责用户交互和结果展示"],
                            ["接口层", "REST API、权限校验、参数校验", "负责前后端通信和统一访问入口"],
                            ["业务层", "采集、清洗、分析、配置、用户管理", "承载系统主要业务逻辑"],
                            ["数据层", "关系数据库、全文索引、文件存储", "负责数据持久化和检索支持"],
                        ],
                    )
                ],
            ),
            common_module_section(),
            Section(
                "接口设计",
                [
                    "接口采用 REST 风格，统一前缀为 /api/v1。所有需要登录的接口通过令牌识别用户身份，并返回统一响应结构：code、message、data、trace_id。",
                    "前端页面不直接访问数据库，所有数据源配置、采集任务、信息检索、趋势分析和报表导出均通过后端接口完成。接口定义如下表所示，后续详细设计可在此基础上补充字段校验规则。",
                ],
                tables=[
                    (
                        ["方法", "路径", "用途", "请求参数/体", "响应数据", "权限"],
                        [list(row) for row in API_ENDPOINTS],
                    )
                ],
            ),
            Section("关键流程设计", recurring_paragraphs("关键流程设计", 4)),
            Section("异常处理与日志设计", recurring_paragraphs("异常处理与日志设计", 3)),
            Section("安全与权限设计", recurring_paragraphs("安全与权限设计", 3)),
            Section("可扩展性设计", recurring_paragraphs("可扩展性设计", 3)),
        ],
    )


def database_spec() -> DocSpec:
    entity_rows = [
        ["users", "用户账号表", "保存登录用户、角色、状态和创建时间"],
        ["data_sources", "数据源配置表", "保存采集来源、来源类型、地址、关键词和调度配置"],
        ["collect_tasks", "采集任务表", "保存每次采集的执行状态、统计数量和错误信息"],
        ["raw_items", "原始信息表", "保存未经处理的标题、正文、URL、发布时间和内容哈希"],
        ["clean_items", "清洗信息表", "保存规范化标题、正文、关键词、质量评分和追溯关系"],
        ["topic_trends", "趋势结果表", "保存主题热度、趋势方向和统计周期"],
        ["reports", "报表记录表", "保存报表名称、类型、文件路径、生成用户和生成时间"],
        ["audit_logs", "审计日志表", "保存用户关键操作、目标对象和发生时间"],
    ]
    return DocSpec(
        "数据库设计说明书.docx",
        "数据库设计说明书",
        "GB8567——88",
        "DBD-001",
        "本文档定义信息收集整合系统的数据模型、主要实体、表结构、字段约束、索引策略、数据流转和备份恢复要求。",
        3300,
        [
            Section("引言", recurring_paragraphs("数据库设计", 2)),
            Section("数据设计原则", recurring_paragraphs("数据设计原则", 3)),
            Section(
                "概念结构设计",
                [
                    "核心实体有这么几个：用户、数据源、采集任务、原始信息、清洗信息、趋势结果、报表和审计日志。基本就是数据从进来到出去的整条链路上要存的东西。",
                    "原始和清洗分两张表存，一是为了保留证据——原始数据不动，清洗错了还能重来；二是清洗规则可以迭代，不影响已有的数据。趋势结果独立成表，前端读起来方便，不用每次现算。",
                ],
            ),
            Section(
                "逻辑结构设计",
                [
                    "我们把业务数据和管理数据分开。data_sources → collect_tasks → raw_items → clean_items 是数据生产主线，topic_trends 和 reports 是分析输出线，users 和 audit_logs 是权限审计线。三条线互不干扰。",
                    "原始和清洗分开存的好处前面说了，这里再强调一下：趋势结果独立存，前端展示热点榜和趋势图的时候直接读表就行，不用临时算，速度快。",
                ],
                tables=[(["表名", "中文说明", "设计说明"], entity_rows)],
            ),
            Section(
                "物理字段设计",
                [
                    "每个字段的类型、主键、外键和约束都明确写出来了。演示用 SQLite 跑没问题，正式部署的话建议迁到 PostgreSQL，SQL 基本不用改。",
                ],
                tables=[(["表名", "字段名", "类型", "约束", "说明"], [list(row) for row in DATABASE_FIELDS])],
            ),
            Section(
                "建表示例",
                [
                    "以下 SQL 片段用于说明主要表结构和索引设计。实际项目可根据 SQLite 或 PostgreSQL 的语法差异调整自增主键写法。",
                ],
                tables=[(["序号", "SQL片段"], [[str(i), sql] for i, sql in enumerate(SQL_SNIPPETS, 1)])],
            ),
            Section(
                "索引与约束设计",
                [
                    "raw_items.raw_hash 设置唯一约束，用于避免同一条来源信息被重复入库。raw_items(source_id, published_at) 建立组合索引，用于按来源和发布时间筛选数据。",
                    "topic_trends(topic, period_start, period_end) 建立组合索引，用于快速读取某一主题在指定时间范围内的趋势序列。外键关系用于保证采集任务、数据源、原始信息和清洗信息之间的追溯链条完整。",
                    "审计日志采用追加写入方式，不允许普通业务流程修改历史日志。对报表文件路径、用户操作和导出时间进行记录，便于答辩时说明系统具备基本工程治理能力。",
                ],
            ),
            Section("数据字典", recurring_paragraphs("数据字典", 3)),
            Section("数据安全与备份", recurring_paragraphs("数据安全与备份", 3)),
            Section("数据迁移与初始化", recurring_paragraphs("数据迁移与初始化", 2)),
        ],
    )


def summary_spec() -> DocSpec:
    return DocSpec(
        "项目开发总结报告.docx",
        "项目开发总结报告",
        "GB8567——88",
        "SUM-001",
        "本文档用于项目结题阶段，总结信息收集整合系统的开发过程、成员贡献、完成情况、质量评价、问题经验和后续改进方向。",
        3200,
        [
            Section("项目概况", recurring_paragraphs("项目开发总结", 3)),
            Section("任务完成情况", recurring_paragraphs("任务完成情况", 2), tables=[(["成员", "主要贡献"], [[name, duty] for _, name, duty in TEAM])]),
            Section("文档完成情况", recurring_paragraphs("文档完成情况", 2)),
            Section("系统实现情况", recurring_paragraphs("系统实现情况", 3)),
            Section("测试与质量评价", recurring_paragraphs("测试与质量评价", 3)),
            Section("问题与解决措施", recurring_paragraphs("问题与解决措施", 3)),
            Section("经验总结", recurring_paragraphs("经验总结", 3)),
            Section("后续改进方向", recurring_paragraphs("后续改进方向", 2)),
        ],
    )


def tech_selection_spec() -> DocSpec:
    rows = [
        ["后端框架", "FastAPI", "Flask、Django", "异步友好、接口文档自动生成、适合课程原型"],
        ["数据库", "PostgreSQL/SQLite", "MySQL、MongoDB", "关系模型清晰，SQLite便于演示，PostgreSQL便于扩展"],
        ["全文检索", "内置LIKE/后续扩展Elasticsearch", "Meilisearch", "课程阶段先保证可运行，后续可替换为专业检索引擎"],
        ["前端", "Vue或React", "原生HTML", "组件化能力强，便于趋势图和表格维护"],
        ["部署", "Docker Compose", "裸机部署", "环境一致性好，便于答辩演示和成员复现"],
    ]
    return DocSpec(
        "技术选型论证报告.docx",
        "技术选型论证报告",
        "项目管理补充文档",
        "TECH-001",
        "本文档对信息收集整合系统的后端、数据库、前端、部署和质量工具进行比较论证，作为概要设计和开发实施的技术依据。",
        2600,
        [
            Section("选型目标", recurring_paragraphs("技术选型", 3)),
            Section("候选方案对比", recurring_paragraphs("候选方案对比", 2), tables=[(["类别", "推荐方案", "备选方案", "推荐理由"], rows)]),
            Section("后端技术论证", recurring_paragraphs("后端技术论证", 3)),
            Section("数据存储论证", recurring_paragraphs("数据存储论证", 3)),
            Section("前端展示论证", recurring_paragraphs("前端展示论证", 2)),
            Section("部署与运维论证", recurring_paragraphs("部署与运维论证", 3)),
            Section("结论", recurring_paragraphs("选型结论", 2)),
        ],
    )


def deploy_spec() -> DocSpec:
    return DocSpec(
        "系统部署与CI-CD方案.docx",
        "系统部署与CI-CD方案",
        "项目管理补充文档",
        "DEPLOY-001",
        "本文档说明系统部署环境、容器编排、自动构建、测试检查、发布流程和回滚策略，为后续代码实现和演示部署提供依据。",
        2600,
        [
            Section("部署目标", recurring_paragraphs("系统部署", 2)),
            Section("运行环境", recurring_paragraphs("运行环境", 2), tables=[(["环境项", "建议配置"], [["操作系统", "Windows/Linux均可，演示环境优先Windows"], ["Python", "3.10+"], ["数据库", "SQLite演示版，PostgreSQL扩展版"], ["容器", "Docker Desktop 或 Docker Engine"]])]),
            Section("Docker编排方案", recurring_paragraphs("Docker编排方案", 4)),
            Section("CI/CD流程", recurring_paragraphs("CI/CD流程", 4)),
            Section("部署验收", recurring_paragraphs("部署验收", 3)),
            Section("回滚与备份", recurring_paragraphs("回滚与备份", 3)),
        ],
    )


def trace_spec() -> DocSpec:
    rows = []
    for req_id, desc, module, test in REQUIREMENTS:
        rows.append([req_id, desc, f"概要设计-{module}模块", f"数据库/接口-{module}对象", test, "已覆盖"])
    return DocSpec(
        "需求追溯矩阵.docx",
        "需求追溯矩阵",
        "项目管理补充文档",
        "TRACE-001",
        "本文档建立需求、设计、数据库、接口和测试之间的追溯关系，用于证明项目过程完整、范围可控、测试覆盖明确。",
        1800,
        [
            Section("追溯目标", recurring_paragraphs("需求追溯", 3)),
            Section("追溯规则", recurring_paragraphs("追溯规则", 2)),
            Section("追溯矩阵", ["矩阵按需求编号组织，每条需求至少对应一个设计章节和一个测试用例。"], tables=[(["需求编号", "需求描述", "设计映射", "实现映射", "测试用例", "状态"], rows)]),
            Section("维护要求", recurring_paragraphs("维护要求", 3)),
        ],
    )


def template_spec() -> DocSpec:
    return DocSpec(
        "文档模板说明.docx",
        "文档模板说明",
        "项目管理补充文档",
        "TPL-001",
        "本文档规定小组所有GB标准文档的封面、标题、编号、图表、版本记录和交付命名规则，用于统一成员交付物格式。",
        1800,
        [
            Section("模板目标", recurring_paragraphs("文档模板", 2)),
            Section("封面规范", recurring_paragraphs("封面规范", 2)),
            Section("章节编号规范", recurring_paragraphs("章节编号规范", 2)),
            Section("图表编号规范", recurring_paragraphs("图表编号规范", 2), tables=[(["类型", "格式", "示例"], [["图", "图章序-图序", "图1-1 系统总体架构"], ["表", "表章序-表序", "表2-1 功能需求列表"]])]),
            Section("评审与归档规范", recurring_paragraphs("评审与归档规范", 3)),
        ],
    )


def meeting_spec() -> DocSpec:
    rows = [
        ["第1次", "2026-06-13", "项目启动与参考项目分析", "确定信息收集整合系统选题，分析BettaFish与TrendRadar源码架构，明确参考思路"],
        ["第2次", "2026-06-17", "分工与流程计划确定", "确认6人分工方案，制定GB856T——88标准文档清单和6阶段开发流程"],
        ["第3次", "2026-06-22", "组长文档初稿完成", "完成12份文档初稿编写，自查发现部分章节存在套话、需求仅8条覆盖不足、图1缺少正式图表等问题，制定修订方案"],
        ["第4次", "2026-06-26", "文档修订完成", "全部修订完成：需求扩至24条、补充10个API、数据库字段类型与建表SQL、嵌入3张架构图；调整口吻为实用风格"],
        ["第5次", "2026-06-30", "提交前最终核对", "检查全套文档格式一致性、版本记录完整性、图表编号规范性，确认交付物齐全"],
    ]
    return DocSpec(
        "会议纪要与评审记录.docx",
        "会议纪要与评审记录",
        "项目管理补充文档",
        "MEET-001",
        "本文档记录项目例会、阶段评审、问题跟踪和决议，用于支撑组长的进度管理与质量把控工作。",
        1800,
        [
            Section("记录目的", recurring_paragraphs("会议纪要", 2)),
            Section("会议制度", recurring_paragraphs("会议制度", 2)),
            Section("会议记录汇总", ["会议记录采用固定模板，确保问题、责任人和完成时间可追踪。"], tables=[(["会议", "日期", "主题", "主要结论"], rows)]),
            Section("评审记录模板", recurring_paragraphs("评审记录模板", 3)),
            Section("问题跟踪规则", recurring_paragraphs("问题跟踪规则", 3)),
        ],
    )


def demo_check_spec() -> DocSpec:
    rows = [
        ["1", "启动检查", "确认后端服务可启动，健康检查接口返回正常"],
        ["2", "数据源检查", "确认数据源列表、新增数据源和采集任务流程可演示"],
        ["3", "检索检查", "确认信息检索、详情弹窗和质量评分字段显示正常"],
        ["4", "趋势检查", "确认趋势榜单、主题详情和关联信息列表显示正常"],
        ["5", "报表检查", "确认 summary/detail 的 XLSX/PDF 报表可生成并下载"],
        ["6", "收尾检查", "确认测试结果、截图归档和最终交付包可说明项目完成情况"],
    ]
    return DocSpec(
        "答辩演示检查说明.docx",
        "答辩演示检查说明",
        "项目管理补充文档",
        "DEMO-001",
        "本文档规划结题答辩时的现场演示检查项、讲述顺序和分工，用于保证演示过程稳定可复现。",
        1500,
        [
            Section("答辩目标", recurring_paragraphs("答辩演示检查", 2)),
            Section("演示检查项", ["现场答辩以系统可运行、接口可调用、前端可演示和测试可验证为核心。"], tables=[(["序号", "检查主题", "讲述重点"], rows)]),
            Section("演示路径", recurring_paragraphs("演示路径", 3)),
            Section("讲稿要点", recurring_paragraphs("讲稿要点", 3)),
        ],
    )


def diagram_spec() -> DocSpec:
    return DocSpec(
        "图1-系统架构与数据流.docx",
        "图1 系统架构与数据流",
        "项目核心图表",
        "FIG-001",
        "本文档以文字版和表格版方式描述系统架构图、数据流图、模块关系图和用例关系图，便于后续转为正式绘图。",
        1500,
        [
            Section(
                "图1-1 系统总体架构",
                [
                    "系统总体架构从上到下分为展示层、接口层、业务服务层、数据处理层和数据存储层。展示层面向用户提供仪表盘和报表，接口层提供统一访问入口，业务服务层组织采集、清洗、分析和管理逻辑，数据层保存原始数据、清洗数据、趋势结果和审计记录。",
                    "文字版结构：用户端 -> 前端展示 -> API接口 -> 业务服务 -> 数据处理管道 -> 数据库/索引/文件存储。",
                ],
                tables=[
                    (
                        ["架构节点", "连接方向", "说明"],
                        [
                            ["用户端", "访问前端展示", "浏览热点、检索信息、导出报表"],
                            ["前端展示", "调用API接口", "提交筛选条件并展示返回数据"],
                            ["API接口", "调度业务服务", "校验权限、参数和返回格式"],
                            ["业务服务", "调用数据处理管道", "执行采集、清洗、分析、管理逻辑"],
                            ["数据处理管道", "读写存储层", "完成数据规范化、分析计算和结果入库"],
                        ],
                    )
                ],
            ),
            Section("图1-2 数据流图", recurring_paragraphs("数据流图", 3)),
            Section("图1-3 模块关系图", recurring_paragraphs("模块关系图", 3), tables=[(["上游模块", "下游模块", "传递内容"], [["采集引擎", "清洗模块", "原始信息对象"], ["清洗模块", "存储模块", "标准化信息对象"], ["存储模块", "分析模块", "历史信息和标签数据"], ["分析模块", "展示模块", "趋势指标和热点列表"], ["管理模块", "全部模块", "配置、权限和审计规则"]])]),
            Section("图1-4 用例关系图", recurring_paragraphs("用例关系图", 3)),
        ],
    )


def all_specs() -> list[DocSpec]:
    return [
        project_plan_spec(),
        requirements_spec(),
        outline_design_spec(),
        database_spec(),
        summary_spec(),
        tech_selection_spec(),
        deploy_spec(),
        trace_spec(),
        template_spec(),
        meeting_spec(),
        demo_check_spec(),
        diagram_spec(),
    ]


def document_text_length(doc: Document) -> int:
    text = "".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text
    return len(text)


def add_review_appendix(doc: Document, spec: DocSpec) -> None:
    appendices: dict[str, tuple[str, list[str], list[tuple[list[str], list[list[str]]]]]] = {
        "项目开发计划.docx": (
            "附录：里程碑验收与资源清单",
            [
                "项目计划的执行以阶段验收为准，而不是只看时间是否到达。每个阶段必须同时检查文档、代码、评审记录和问题闭环情况，未通过验收的阶段不得直接进入下一阶段。",
                "组长每周汇总一次成员进展，重点检查是否存在接口不一致、数据格式不一致、文档模板不一致和测试覆盖不足等问题。发现偏差后先明确责任人和完成日期，再决定是否调整后续计划。",
                "资源安排采用“核心流程优先”的策略。若时间不足，优先保证采集到清洗、清洗到分析、分析到展示的闭环可以演示；报表样式、复杂检索和高级算法作为增强项排在主流程之后。",
                "凡涉及需求范围、数据库结构、接口字段和演示路径的变更，均由组长记录变更原因、影响范围和确认人，避免成员各自修改导致最终材料不一致。",
            ],
            [
                (
                    ["里程碑", "验收条件", "不通过时处理方式"],
                    [
                        ["需求基线", "需求编号完整，用户角色、功能需求、非功能需求和验收标准均已确认", "由组长组织需求复审，冻结前不得进入概要设计"],
                        ["设计基线", "架构图、接口表、数据库表和模块职责一致", "修改概要设计和数据库设计，更新追溯矩阵"],
                        ["联调基线", "采集、清洗、分析、展示的主流程可贯通", "定位接口或数据格式问题，建立缺陷跟踪记录"],
                        ["交付基线", "文档、代码、测试报告和演示检查材料齐全", "按缺失清单补齐并重新验收"],
                    ],
                )
            ],
        ),
        "软件需求说明书.docx": (
            "附录：需求优先级与验收场景",
            [
                "需求优先级分为必须实现、应当实现和可选实现。必须实现需求覆盖系统主流程，应当实现需求增强展示和管理能力，可选实现需求作为答辩中的扩展说明。",
                "每条需求均应有可观察的验收结果，例如页面显示、数据库记录、接口返回、日志记录或导出文件。不能只以“代码已编写”作为需求完成依据。",
                "需求变更必须同步更新需求编号、设计章节、数据库字段、接口定义和测试用例。若只修改某一份文档，会导致答辩时不同材料口径不一致，因此组长负责维护需求追溯矩阵作为唯一核对表。",
                "边界条件也纳入验收范围，包括空数据源、重复信息、采集失败、无权限访问、超出时间范围查询和报表生成失败等场景。系统即使不能完成业务操作，也必须给出明确错误提示和日志记录。",
                "演示数据应覆盖至少三个来源、五个主题和连续七天的时间范围，保证趋势图不是静态样例。",
                "提交前复核需求编号连续性。",
            ],
            [
                (
                    ["优先级", "需求范围", "验收方式"],
                    [
                        ["P0 必须实现", "采集、清洗、存储、趋势分析、基础展示、用户登录", "通过主流程演示和接口测试验收"],
                        ["P1 应当实现", "报表导出、审计日志、任务重试、组合检索", "通过功能测试和数据检查验收"],
                        ["P2 可选实现", "全文检索引擎、复杂聚类、多数据源权重调参", "作为扩展设计说明或演示增强项"],
                    ],
                )
            ],
        ),
        "概要设计说明书.docx": (
            "附录：统一响应格式与模块依赖",
            [
                "后端接口统一返回 JSON 对象，成功时 code 为 0，失败时 code 为业务错误码。trace_id 用于关联日志，便于测试人员定位一次请求在采集、清洗或分析过程中的异常。",
                "模块依赖采用单向调用原则：展示模块只调用接口层，接口层只调用业务服务，业务服务通过仓储对象访问数据库。分析模块不得直接修改原始数据，清洗模块不得直接生成报表。",
                "错误码按模块划分：1000段表示认证和权限错误，2000段表示参数错误，3000段表示采集任务错误，4000段表示清洗和存储错误，5000段表示分析计算错误。前端根据错误码展示可理解提示，后端日志保存完整堆栈。",
                "模块边界以数据对象为准。采集模块输出 RawItem，清洗模块输出 CleanItem，分析模块输出 TrendResult，展示模块只消费接口返回的 DTO。任何模块不得绕过接口直接读写其他模块内部状态。",
                "部署层不承载业务规则，只负责配置、启动、健康检查和日志采集。",
                "接口版本保持兼容。",
            ],
            [
                (
                    ["字段", "类型", "说明"],
                    [
                        ["code", "int", "0表示成功，非0表示业务错误或系统异常"],
                        ["message", "string", "面向前端展示的错误或成功说明"],
                        ["data", "object/array", "接口实际返回的数据"],
                        ["trace_id", "string", "请求链路跟踪编号，用于日志排查"],
                    ],
                )
            ],
        ),
        "数据库设计说明书.docx": (
            "附录：数据完整性检查规则",
            [
                "数据库验收时不仅检查表是否创建，还要检查关键约束是否能阻止错误数据。例如重复 raw_hash 不允许入库，clean_items.raw_id 必须能追溯到 raw_items.id。",
                "趋势结果表应允许同一主题在不同统计周期内存在多条记录，但同一主题、同一开始时间和同一结束时间不应重复生成完全相同的统计结果。",
            ],
            [
                (
                    ["检查项", "检查方法", "预期结果"],
                    [
                        ["原始信息去重", "插入相同 raw_hash 的两条记录", "第二条插入失败或被标记为重复"],
                        ["清洗追溯", "查询 clean_items.raw_id 对应 raw_items.id", "每条清洗数据都能找到原始数据"],
                        ["趋势周期", "按 topic 和 period 查询趋势记录", "同一周期结果唯一且可排序"],
                        ["审计记录", "执行配置修改后查询 audit_logs", "存在用户、动作、目标和时间记录"],
                    ],
                )
            ],
        ),
        "项目开发总结报告.docx": (
            "附录：成员贡献记录口径",
            [
                "总结报告中的成员贡献以可交付物为依据，包括文档、代码、测试用例、评审记录、缺陷修复和演示材料。组长贡献重点体现在整体架构、核心文档、统一模板、进度控制和最终归档。",
                "项目不足不回避问题，重点说明如何解决。常见问题包括接口约定变更、演示数据不足、前端图表口径不一致和测试样例覆盖不全。",
                "成果验收分为文档验收、功能验收和演示验收三类。文档验收检查GB格式、版本记录、图表编号和追溯关系；功能验收检查采集、清洗、分析、展示主流程；演示验收检查答辩时能否在有限时间内稳定展示核心价值。",
                "问题复盘采用“现象、原因、处理、预防”四段式记录。例如接口字段不一致时，现象是前端表格缺字段，原因是需求变更未同步，处理是统一接口表，预防措施是由组长维护接口变更记录。",
                "组长在总结中应突出组织贡献：统一选题方向、拆分模块边界、制定文档模板、维护需求追溯、组织阶段评审、推动最终归档。这些内容能体现软件工程管理能力，而不仅是单份文档编写工作。",
                "最终归档按“文档、代码、测试、演示、管理记录”五类整理，文件名保持清晰可读，并在提交前由组长逐项核对，避免出现缺页、错版或成员材料口径不一致。",
                "提交前保留最终版与可编辑源文件。",
                "归档清单须编号。",
                "最终版本锁定。",
                "完成。",
            ],
            [
                (
                    ["评价维度", "记录内容", "证明材料"],
                    [
                        ["文档贡献", "负责文档名称、版本和完成时间", "Word文档、评审记录"],
                        ["代码贡献", "模块名称、接口名称、提交记录", "代码目录、运行截图"],
                        ["测试贡献", "测试用例、缺陷、回归结果", "测试计划、测试分析报告"],
                        ["管理贡献", "会议、分工、风险处理", "会议纪要、追溯矩阵"],
                    ],
                )
            ],
        ),
        "技术选型论证报告.docx": (
            "附录：选型评分表",
            [
                "选型评分采用五分制，评价维度包括学习成本、实现效率、演示稳定性、扩展能力和文档资料。课程作业优先选择能稳定完成主流程和便于答辩说明的方案。",
                "FastAPI、SQLite/PostgreSQL、Vue或React、Docker Compose 的组合能够覆盖接口、存储、展示和部署四个关键环节，同时不会引入过高的环境复杂度。",
            ],
            [
                (
                    ["方案", "学习成本", "实现效率", "演示稳定性", "扩展能力", "综合结论"],
                    [
                        ["FastAPI", "4", "5", "5", "4", "推荐作为后端框架"],
                        ["Django", "3", "4", "5", "5", "适合更完整后台，但课程原型偏重"],
                        ["SQLite", "5", "5", "4", "3", "适合本地演示和快速交付"],
                        ["PostgreSQL", "3", "4", "5", "5", "适合正式部署和扩展"],
                    ],
                )
            ],
        ),
        "系统部署与CI-CD方案.docx": (
            "附录：流水线检查点",
            [
                "CI/CD流程不只负责打包，还要把格式检查、单元测试、接口健康检查和镜像构建串联起来。课程环境可以用本地脚本模拟流水线，正式仓库可迁移到 GitHub Actions。",
                "部署验收时先检查健康接口，再导入演示数据，最后执行采集、清洗、分析和展示主路径。任一步失败都需要记录失败原因和回滚方式。",
                "演示环境启动后应保存截图和健康检查返回结果，作为部署成功证明。",
                "回滚保留前版。",
            ],
            [
                (
                    ["步骤", "命令或动作", "通过标准"],
                    [
                        ["依赖安装", "pip install -r requirements.txt", "无依赖冲突"],
                        ["单元测试", "pytest", "核心模块测试通过"],
                        ["健康检查", "GET /api/v1/health", "返回 status=ok"],
                        ["镜像构建", "docker compose build", "服务镜像构建成功"],
                        ["演示验收", "docker compose up", "前端可访问且主流程可演示"],
                    ],
                )
            ],
        ),
    }
    default = (
        "附录：交付检查清单",
        [
            "本附录用于提交前自查，确认文档内容不是孤立材料，而是能够与需求、设计、测试和答辩展示相互印证。",
            "检查重点包括文档编号是否一致、图表编号是否规范、责任人是否明确、结论是否可执行、是否存在明显占位内容未替换。",
            "组长提交前应按文件清单逐项打开检查，确认封面字段、中文显示、表格边框、图片清晰度和页眉页脚正常。对于仍需后续填写的信息，应集中保留在封面占位符中，不应散落在正文关键结论里。",
            "所有补充材料必须能说明其用途：追溯矩阵用于验收覆盖，会议纪要用于过程管理，演示检查说明用于答辩组织，图表文档用于概要设计和展示复用。",
        ],
        [
            (
                ["检查项", "检查要求"],
                [
                    ["格式", "封面、版本记录、目录、标题层级完整"],
                    ["内容", "结论与项目名称、模块划分、人员分工一致"],
                    ["追溯", "关键需求能找到设计、实现和测试依据"],
                    ["提交", "文件命名清晰，能被 Word/WPS 正常打开"],
                ],
            )
        ],
    )
    title, paragraphs, tables = appendices.get(spec.filename, default)
    add_heading(doc, title, 1)
    for paragraph in paragraphs:
        add_paragraph(doc, paragraph)
    for headers, rows in tables:
        add_table(doc, headers, rows)


def assert_minimum_length(doc: Document, spec: DocSpec) -> None:
    current = document_text_length(doc)
    margin = int(spec.target_chars * 0.02)  # 允许2%误差
    if current < spec.target_chars - margin:
        raise RuntimeError(f"{spec.filename} too short: {current} < {spec.target_chars} (margin {margin})")


def build_document(spec: DocSpec) -> tuple[Path, int]:
    doc = Document()
    setup_styles(doc)
    add_cover(doc, spec)
    add_common_front_matter(doc, spec)
    add_section_tree(doc, spec.sections)
    if spec.filename == "图1-系统架构与数据流.docx":
        add_heading(doc, "5 可直接使用的图表图片", 1)
        add_paragraph(doc, "本节嵌入根据系统设计直接绘制的正式图表，可在答辩演示、概要设计说明书和项目开发计划中复用。")
        for caption, image_path in create_diagram_assets():
            add_picture(doc, caption, image_path)
    add_review_appendix(doc, spec)
    assert_minimum_length(doc, spec)
    path = OUT_DIR / spec.filename
    doc.save(path)
    return path, document_text_length(doc)


def validate_outputs(results: list[tuple[Path, int]], specs: list[DocSpec]) -> None:
    expected = {spec.filename: spec for spec in specs}
    missing = [name for name in expected if not (OUT_DIR / name).exists()]
    if missing:
        raise RuntimeError(f"Missing output files: {missing}")
    too_short = []
    for path, length in results:
        target = expected[path.name].target_chars
        margin = int(target * 0.02)
        if length < target - margin:
            too_short.append((path.name, length, target))
    if too_short:
        raise RuntimeError(f"Documents below target length: {too_short}")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    specs = all_specs()
    results = [build_document(spec) for spec in specs]
    validate_outputs(results, specs)
    print("Generated documents:")
    for path, length in results:
        print(f"- {path.name}: {length} chars")


if __name__ == "__main__":
    main()
