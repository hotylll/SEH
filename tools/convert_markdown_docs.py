from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
PROJECT_NAME = "集思 · 信息收集整合系统"
PROJECT_ENGLISH_NAME = "Jisi InfoHub"
EXCLUDED_PARTS = {
    ".git",
    ".venv",
    "venv",
    "__pycache__",
    "GB856T——88",
    "集思 · 信息收集整合系统-最终交付包",
    "集思 · 信息收集整合系统-最终交付包",
}


def set_run_font(run, size: int = 11, bold: bool = False, font: str = "宋体") -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_paragraph_font(paragraph, size: int = 11, bold: bool = False, font: str = "宋体") -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, font=font)


def add_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    text = text.replace("`", "")
    text = text.replace("**", "")
    text = text.replace("__", "")
    text = text.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    return text.strip()


def is_separator_row(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def split_table_row(line: str) -> list[str]:
    return [clean_inline(cell.strip()) for cell in line.strip().strip("|").split("|")]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index in range(width):
            text = row[col_index] if col_index < len(row) else ""
            cells[col_index].text = text
            for paragraph in cells[col_index].paragraphs:
                set_paragraph_font(paragraph, size=10, bold=row_index == 0)
            if row_index == 0:
                add_shading(cells[col_index], "D9EAF7")
    doc.add_paragraph()


def markdown_files() -> list[Path]:
    files: list[Path] = []
    for path in ROOT.rglob("*.md"):
        rel_parts = set(path.relative_to(ROOT).parts)
        if EXCLUDED_PARTS.intersection(rel_parts):
            continue
        files.append(path)
    return sorted(files, key=lambda item: item.as_posix())


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for name, size in [("Title", 18), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "黑体"
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_cover(doc: Document, source: Path, title: str) -> None:
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(PROJECT_NAME)
    set_run_font(run, size=20, bold=True, font="黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(PROJECT_ENGLISH_NAME)
    set_run_font(run, size=13, bold=False, font="Times New Roman")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, size=18, bold=True, font="黑体")
    doc.add_page_break()


def convert_file(path: Path) -> Path:
    text = path.read_text(encoding="utf-8-sig")
    lines = text.splitlines()
    title = clean_inline(next((line.lstrip("#").strip() for line in lines if line.startswith("#")), path.stem))

    doc = Document()
    setup_document(doc)
    add_cover(doc, path, title)

    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run("\n".join(code_lines))
                set_run_font(run, size=9, font="Consolas")
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped or stripped == "---":
            i += 1
            continue

        if "|" in stripped and i + 1 < len(lines) and is_separator_row(lines[i + 1]):
            rows = [split_table_row(stripped)]
            i += 2
            while i < len(lines) and "|" in lines[i].strip() and lines[i].strip():
                rows.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 3)
            doc.add_heading(clean_inline(heading.group(2)), level=level)
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        quote = re.match(r"^>\s*(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(clean_inline(bullet.group(1)))
            set_paragraph_font(p)
        elif numbered:
            p = doc.add_paragraph(style="List Number")
            p.add_run(clean_inline(numbered.group(1)))
            set_paragraph_font(p)
        elif quote:
            p = doc.add_paragraph(clean_inline(quote.group(1)))
            p.paragraph_format.left_indent = Pt(18)
            set_paragraph_font(p, size=10)
        else:
            p = doc.add_paragraph(clean_inline(stripped))
            set_paragraph_font(p)
        i += 1

    out = path.with_suffix(".docx")
    doc.save(out)
    return out


def main() -> None:
    outputs = [convert_file(path) for path in markdown_files()]
    for output in outputs:
        print(output.relative_to(ROOT).as_posix())
    print(f"converted {len(outputs)} markdown files")


if __name__ == "__main__":
    main()
